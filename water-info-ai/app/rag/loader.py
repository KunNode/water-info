"""Document loaders for the knowledge base."""

from __future__ import annotations

import io
import mimetypes
import re
from pathlib import Path

from app.rag.models import LoadedDocument, TextBlock


def detect_mime(filename: str, provided_mime: str | None = None) -> str:
    if provided_mime and provided_mime != "application/octet-stream":
        return provided_mime
    guessed, _ = mimetypes.guess_type(filename)
    return guessed or "text/plain"


def _normalize_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def _guess_title(filename: str, explicit_title: str | None = None) -> str:
    if explicit_title:
        return explicit_title.strip()
    return Path(filename).stem or "未命名文档"


def _flush_paragraph(
    blocks: list[TextBlock],
    paragraph_lines: list[str],
    heading_path: list[str],
    metadata: dict | None = None,
) -> None:
    text = "\n".join(line.rstrip() for line in paragraph_lines).strip()
    if text:
        blocks.append(TextBlock(text=text, heading_path=list(heading_path), metadata=metadata or {}))
    paragraph_lines.clear()


def _load_markdown_text(raw_text: str, *, title: str, mime: str) -> LoadedDocument:
    lines = _normalize_text(raw_text).split("\n")
    heading_stack: list[str] = []
    paragraph_lines: list[str] = []
    blocks: list[TextBlock] = []

    for line in lines:
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if heading_match:
            _flush_paragraph(blocks, paragraph_lines, heading_stack)
            level = len(heading_match.group(1))
            heading = heading_match.group(2).strip()
            heading_stack = heading_stack[: level - 1]
            heading_stack.append(heading)
            continue
        if not line.strip():
            _flush_paragraph(blocks, paragraph_lines, heading_stack)
            continue
        paragraph_lines.append(line)

    _flush_paragraph(blocks, paragraph_lines, heading_stack)

    if not blocks and raw_text.strip():
        blocks.append(TextBlock(text=_normalize_text(raw_text)))
    return LoadedDocument(title=title, mime=mime, raw_text=_normalize_text(raw_text), blocks=blocks)


def _load_plain_text(raw_text: str, *, title: str, mime: str) -> LoadedDocument:
    normalized = _normalize_text(raw_text)
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", normalized) if part.strip()]
    blocks = [TextBlock(text=paragraph) for paragraph in paragraphs] or [TextBlock(text=normalized)]
    return LoadedDocument(title=title, mime=mime, raw_text=normalized, blocks=blocks)


def _load_docx(content: bytes, *, title: str, mime: str) -> LoadedDocument:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - exercised only when dependency missing
        raise RuntimeError("python-docx is required for DOCX ingestion") from exc

    document = Document(io.BytesIO(content))
    heading_stack: list[str] = []
    blocks: list[TextBlock] = []
    raw_parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        raw_parts.append(text)
        style_name = getattr(paragraph.style, "name", "")
        heading_match = re.match(r"Heading\s+(\d+)", style_name or "")
        if heading_match:
            level = int(heading_match.group(1))
            heading_stack = heading_stack[: level - 1]
            heading_stack.append(text)
            continue
        blocks.append(TextBlock(text=text, heading_path=list(heading_stack)))

    raw_text = "\n\n".join(raw_parts)
    return LoadedDocument(title=title, mime=mime, raw_text=_normalize_text(raw_text), blocks=blocks)


def _load_pdf(content: bytes, *, title: str, mime: str) -> LoadedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - exercised only when dependency missing
        raise RuntimeError("pypdf is required for PDF ingestion") from exc

    reader = PdfReader(io.BytesIO(content))
    blocks: list[TextBlock] = []
    raw_pages: list[str] = []

    for page_number, page in enumerate(reader.pages, start=1):
        page_text = _normalize_text(page.extract_text() or "")
        if not page_text:
            continue
        raw_pages.append(page_text)
        for paragraph in re.split(r"\n\s*\n", page_text):
            paragraph = paragraph.strip()
            if paragraph:
                blocks.append(TextBlock(text=paragraph, metadata={"page": page_number}))

    raw_text = "\n\n".join(raw_pages)
    return LoadedDocument(title=title, mime=mime, raw_text=raw_text, blocks=blocks)


def load_uploaded_document(
    filename: str,
    content: bytes,
    *,
    title: str | None = None,
    mime: str | None = None,
) -> LoadedDocument:
    detected_mime = detect_mime(filename, mime)
    document_title = _guess_title(filename, title)
    suffix = Path(filename).suffix.lower()

    if suffix in {".md", ".markdown"} or detected_mime in {"text/markdown", "text/x-markdown"}:
        return _load_markdown_text(content.decode("utf-8", errors="ignore"), title=document_title, mime=detected_mime)
    if suffix == ".txt" or detected_mime.startswith("text/"):
        return _load_plain_text(content.decode("utf-8", errors="ignore"), title=document_title, mime=detected_mime)
    if suffix == ".docx" or detected_mime in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }:
        return _load_docx(content, title=document_title, mime=detected_mime)
    if suffix == ".pdf" or detected_mime == "application/pdf":
        return _load_pdf(content, title=document_title, mime=detected_mime)
    raise RuntimeError(f"不支持的文档类型: {filename}")


def load_plain_document(title: str, raw_text: str, mime: str) -> LoadedDocument:
    if mime in {"text/markdown", "text/x-markdown"}:
        return _load_markdown_text(raw_text, title=title, mime=mime)
    return _load_plain_text(raw_text, title=title, mime=mime)
